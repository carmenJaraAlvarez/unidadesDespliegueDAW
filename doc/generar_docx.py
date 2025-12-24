#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script to generate Word document from Markdown
Usage: python generar_docx.py
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_markdown_to_docx(md_file, docx_file):
    """Parse markdown and create Word document"""
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_table = False
    table_data = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line:
            if not in_table:
                doc.add_paragraph()
            i += 1
            continue
        
        # Heading level 1
        if line.startswith('# ') and not line.startswith('## '):
            heading = doc.add_heading(line[2:], level=1)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(33, 136, 56)
            i += 1
            continue
        
        # Heading level 2
        if line.startswith('## ') and not line.startswith('### '):
            heading = doc.add_heading(line[3:], level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0, 114, 188)
            i += 1
            continue
        
        # Heading level 3
        if line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(33, 136, 56)
            i += 1
            continue
        
        # Horizontal rule - page break
        if line.startswith('---'):
            doc.add_page_break()
            i += 1
            continue
        
        # Table
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            
            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data.append(cells)
            i += 1
            
            # Check if next line is separator or end of table
            if i < len(lines):
                next_line = lines[i].rstrip()
                if next_line.startswith('|') and '---' in next_line:
                    # Skip separator line
                    i += 1
                    continue
                elif not next_line.startswith('|'):
                    # End of table, create it
                    if table_data:
                        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                        table.style = 'Light Grid Accent 1'
                        
                        for row_idx, row_data in enumerate(table_data):
                            for col_idx, cell_data in enumerate(row_data):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_data
                                # Bold header row
                                if row_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True
                    
                    in_table = False
                    table_data = []
            continue
        
        # Bullet list
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            # Check for bold patterns
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, line)
        i += 1
    
    doc.save(docx_file)
    print(f"✅ Documento Word generado: {docx_file}")

def add_formatted_text(paragraph, text):
    """Add text with bold formatting"""
    # Pattern to find bold text **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            # Regular text
            paragraph.add_run(part)

if __name__ == '__main__':
    md_file = 'Propuesta_Didactica_Despliegue_DAW.md'
    docx_file = 'Propuesta_Didactica_Despliegue_DAW.docx'
    
    try:
        parse_markdown_to_docx(md_file, docx_file)
        print("\n📄 Archivo Markdown: " + md_file)
        print("📝 Archivo Word:     " + docx_file)
        print("\n💡 Tip: Edita el archivo .md y vuelve a ejecutar este script para regenerar el .docx")
    except Exception as e:
        print(f"❌ Error: {e}")

